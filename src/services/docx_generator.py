import os
import tempfile
import time
import traceback
import copy
import pathlib
import re

from docx import Document
from docx.shared import Mm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def replace_in_runs(paragraph, datas):
    """ТВОЯ текущая замена (как было) — оставляем для всех бланков кроме СИЗ."""
    if not datas or not isinstance(datas, dict):
        return
    if paragraph is None or not getattr(paragraph, "runs", None):
        return

    runs = paragraph.runs
    if not runs:
        return

    touched = False
    for run in runs:
        t = run.text or ""
        if not t:
            continue
        original = t
        for ph, val in datas.items():
            if ph in t:
                t = t.replace(ph, "" if val is None else str(val))
        if t != original:
            run.text = t
            touched = True

    if touched:
        return

    full_text = "".join((r.text or "") for r in runs)
    if not full_text:
        return
    if not any(ph in full_text for ph in datas.keys()):
        return

    new_text = full_text
    for ph, val in datas.items():
        new_text = new_text.replace(ph, "" if val is None else str(val))

    pos = 0
    for r in runs:
        ln = len(r.text or "")
        r.text = new_text[pos:pos + ln]
        pos += ln

    if runs and pos < len(new_text):
        runs[-1].text = (runs[-1].text or "") + new_text[pos:]


def replace_in_runs_force(paragraph, datas):
    """
    Агрессивная замена: всегда работает даже если плейсхолдер разбит на runs.
    Применять ТОЛЬКО для СИЗ.
    """
    if not datas or not isinstance(datas, dict):
        return
    if paragraph is None or not getattr(paragraph, "runs", None):
        return

    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join((r.text or "") for r in runs)
    if not full_text:
        return
    if not any(ph in full_text for ph in datas.keys()):
        return

    new_text = full_text
    for ph, val in datas.items():
        new_text = new_text.replace(ph, "" if val is None else str(val))

    pos = 0
    for r in runs:
        ln = len(r.text or "")
        r.text = new_text[pos:pos + ln]
        pos += ln

    if runs and pos < len(new_text):
        runs[-1].text = (runs[-1].text or "") + new_text[pos:]


def _table_has_text(table, needle: str) -> bool:
    """Проверяет, встречается ли needle в тексте ячеек таблицы (быстро и надёжно)."""
    try:
        for row in table.rows:
            for cell in row.cells:
                if needle in (cell.text or ""):
                    return True
    except Exception:
        pass
    return False


_MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}


def _parse_user_date(date_str: str) -> tuple | None:
    s = (date_str or "").strip().lower()
    if not s:
        return None

    s = re.sub(r"\s*(г\.?|год\.?)\s*$", "", s, flags=re.IGNORECASE).strip()

    m = re.search(r"(\d{1,2})\D+(\d{1,2})\D+(\d{2,4})", s)
    if m:
        d = int(m.group(1))
        mo = int(m.group(2))
        y = int(m.group(3))
        if y < 100:
            y += 2000
        return d, mo, y

    for mo_num, mo_name in _MONTHS_RU.items():
        if mo_name in s:
            m2 = re.search(r"(\d{1,2})\D+(\d{2,4})", s)
            if m2:
                d = int(m2.group(1))
                y = int(m2.group(2))
                if y < 100:
                    y += 2000
                return d, mo_num, y

    return None


def _fmt_long_ru(date_tuple) -> str:
    if not date_tuple:
        return ""
    d, mo, y = date_tuple
    return f"{d} {_MONTHS_RU.get(mo, '')} {y}"


def _fmt_short_g(date_tuple) -> str:
    if not date_tuple:
        return ""
    d, mo, y = date_tuple
    s = f"{d:02d}.{mo:02d}.{y}г."
    return s.replace(".", "\u2024")


def _set_cell_width_mm(cell, mm: float):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:type"), "dxa")
    tcW.set(qn("w:w"), str(int(mm * 56.7)))


def _widen_osnovanie_left_cell(table, add_mm: float = 3.0):
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        left = row.cells[0]
        right = row.cells[1]

        txt = (left.text or "").replace("\u00A0", " ").strip()
        if "Основание" not in txt:
            continue

        try:
            l_mm = left.width.mm
        except Exception:
            l_mm = None
        try:
            r_mm = right.width.mm
        except Exception:
            r_mm = None

        if l_mm is not None and r_mm is not None:
            _set_cell_width_mm(left, l_mm + add_mm)
            _set_cell_width_mm(right, max(10.0, r_mm - add_mm))
        else:
            _set_cell_width_mm(left, 100.0)
            _set_cell_width_mm(right, 80.0)

        return


def _find_template(filename: str) -> str | None:
    here = pathlib.Path(__file__).resolve().parent
    project_root = here.parent.parent

    candidates = [
        pathlib.Path(filename),
        here / filename,
        project_root / filename,
        project_root / "src" / filename,
        project_root.parent / filename,
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return None


def get_template_paths():
    main_path = _find_template("Main_data1.docx")
    st_path = _find_template("test.docx")
    return main_path, st_path


def _clear_body_keep_sectpr(doc: Document):
    body = doc.element.body
    sectpr = None

    for child in list(body):
        if child.tag == qn("w:sectPr") or child.tag.endswith("}sectPr"):
            sectpr = child
        body.remove(child)

    if sectpr is not None:
        body.append(sectpr)


def _one_empty_line(doc: Document, gap_mm: int = 3):
    p = doc.add_paragraph("")
    pf = p.paragraph_format
    pf.space_before = 0
    pf.space_after = Mm(gap_mm)
    pf.line_spacing = 1
    return p


def _anchor_paragraph(doc: Document):
    p = doc.add_paragraph("")
    pf = p.paragraph_format
    pf.space_before = 0
    pf.space_after = 0
    pf.line_spacing = 1
    return p


def _compact_paragraphs(doc: Document):
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf is None:
            continue
        pf.space_before = 0
        pf.space_after = 0
        pf.line_spacing = 1


def _table_to_inline(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        return
    for child in list(tblPr):
        if child.tag in (qn("w:tblpPr"), qn("w:tblOverlap")):
            tblPr.remove(child)


def _set_tbl_fixed_layout(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        return
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _set_row_cant_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = trPr.find(qn("w:cantSplit"))
    if cant is None:
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def _keep_table_rows_together(table):
    for row in table.rows:
        _set_row_cant_split(row)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.keep_together = True
                pf.keep_with_next = True


def _force_times_new_roman_in_table(table, size_pt: int | None = None):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"

                    rPr = r._element.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = OxmlElement("w:rFonts")
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:ascii"), "Times New Roman")
                    rFonts.set(qn("w:hAnsi"), "Times New Roman")
                    rFonts.set(qn("w:cs"), "Times New Roman")

                    if size_pt is not None:
                        r.font.size = Pt(size_pt)


def _get_count(blanks_count: dict, *keys):
    mx = 0
    for k in keys:
        try:
            mx = max(mx, int(blanks_count.get(k, 0) or 0))
        except Exception:
            pass
    return mx


def _build_plan(blanks_count: dict):
    plan = []

    plan += [("main", 0)] * _get_count(blanks_count, "А", "A")
    plan += [("main", 1)] * _get_count(blanks_count, "Б", "B")
    plan += [("main", 2)] * _get_count(blanks_count, "В", "V")
    plan += [("main", 3)] * _get_count(blanks_count, "ПП", "PP")
    plan += [("main", 4)] * _get_count(blanks_count, "СИЗ", "SIZ", "ОТ", "OT")

    plan += [("st", 0)] * _get_count(blanks_count, "С1", "C1")
    plan += [("st", 1)] * _get_count(blanks_count, "С2", "C2")
    plan += [("st", 2)] * _get_count(blanks_count, "С3", "C3")
    plan += [("st", 3)] * _get_count(blanks_count, "Т1", "T1")
    plan += [("st", 4)] * _get_count(blanks_count, "Т2", "T2")
    plan += [("st", 5)] * _get_count(blanks_count, "Т3", "T3")

    return plan


def _group_for_st_table(table_index: int) -> str:
    if table_index in (0, 3):
        return "1 (первая)"
    if table_index in (1, 4):
        return "2 (вторая)"
    if table_index in (2, 5):
        return "3 (третья)"
    return "2 (вторая)"


def _append_table_like_word(out_doc: Document, src_table):
    t = out_doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    t._tbl.getparent().replace(t._tbl, copy.deepcopy(src_table._tbl))
    return out_doc.tables[-1]


def fill_docx_template(dt):
    if not dt:
        return None

    try:
        full_name = (dt.get("Выдано", "") or "").strip()
        name_parts = [p for p in full_name.split() if p]

        user_date_tuple = _parse_user_date(dt.get("Дата", "") or "")

        try:
            blank_start = int(dt.get("Удост_№", 0) or 0)
        except Exception:
            blank_start = 0

        placeholders_base = {
            "{surname}": name_parts[0] if len(name_parts) > 0 else "",
            "{name}": name_parts[1] if len(name_parts) > 1 else "",
            "{father_name}": name_parts[2] if len(name_parts) > 2 else "",
            "{work_feat}": dt.get("Должность", "") or "",
            "{organization_name}": dt.get("Место работы", "") or "",
            "{organization _name}": dt.get("Место работы", "") or "",
            "{protocol_№}": str(dt.get("ПРТ_№", "") or ""),
            "{protocol}": str(dt.get("ПРТ_№", "") or ""),
            "{data}": _fmt_long_ru(user_date_tuple),
            "{data_start}": "12.10.2025",
            "{data_end}": "12.10.2028",
        }

        blanks_count = dt.get("blanks_count", {}) or {}
        plan = _build_plan(blanks_count)
        if not plan:
            return None

        main_path, st_path = get_template_paths()
        if not main_path:
            print("❌ Не найден Main_data1.docx")
            return None
        if not st_path:
            print("❌ Не найден test.docx")
            return None

        main_doc = Document(main_path)
        st_doc = Document(st_path)

        out_doc = Document(main_path)
        _clear_body_keep_sectpr(out_doc)

        for idx, (src, table_idx) in enumerate(plan):
            if idx > 0:
                _one_empty_line(out_doc, gap_mm=3)
                _anchor_paragraph(out_doc)

            if src == "main":
                if len(main_doc.tables) <= table_idx:
                    raise ValueError(f"Main_data1.docx: нет таблицы {table_idx}, всего {len(main_doc.tables)}")
                src_table = main_doc.tables[table_idx]
            else:
                if len(st_doc.tables) <= table_idx:
                    raise ValueError(f"test.docx: нет таблицы {table_idx}, всего {len(st_doc.tables)}")
                src_table = st_doc.tables[table_idx]

            inserted_table = _append_table_like_word(out_doc, src_table)

            _table_to_inline(inserted_table)
            _set_tbl_fixed_layout(inserted_table)
            _keep_table_rows_together(inserted_table)

            if src == "st":
                _force_times_new_roman_in_table(inserted_table, size_pt=None)
                if table_idx in (0, 1, 2):
                    _widen_osnovanie_left_cell(inserted_table, add_mm=3.0)

        for i, table in enumerate(out_doc.tables):
            # plan[i] оставляем как было (для логики дат/групп), но СИЗ определяем по тексту таблицы
            src, src_table_index = plan[i]

            ph = dict(placeholders_base)
            ph["{blank_№}"] = str(blank_start + i)

            if src == "st":
                ph["{rang_group}"] = _group_for_st_table(src_table_index)
                if src_table_index in (0, 1, 2):
                    ph["{data}"] = _fmt_short_g(user_date_tuple)

            # === ВОТ КЛЮЧЕВАЯ ПРАВКА ===
            is_siz = _table_has_text(table, "{protocol_№}")
            # ===========================

            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if is_siz:
                            replace_in_runs_force(paragraph, ph)  # только СИЗ
                        else:
                            replace_in_runs(paragraph, ph)        # все остальные

        _compact_paragraphs(out_doc)

        temp_dir = tempfile.gettempdir()
        ts = int(time.time())
        out_path = os.path.join(temp_dir, f"blanks_{ts}.docx")
        out_doc.save(out_path)

        dt["Удост_№"] = str(blank_start + len(out_doc.tables))
        return out_path

    except Exception as e:
        print("💥 ОШИБКА:", type(e).__name__, str(e))
        print(traceback.format_exc())
        return None


def filldocxtemplate(dt):
    return fill_docx_template(dt)
